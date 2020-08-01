import cv2 
import argparse
import sys
import numpy as np
import os.path
import glob
import os
from docx import Document
from docx.shared import Inches
import time
class No_Plate_Detector:
	def __init__(self):
		self.confThreshold = 0.1 
		self.nmsThreshold = 0.1 
		self.inpWidth = 416  
		self.inpHeight = 416 

		self.classes  = r'API\classes\classes.names'
		# self.classes  = r'classes\classes.names'
		# with open(classesfile1, 'rt') as f:
		# 	classes1 = f.read().rstrip('\n').split('\n')
		# self.classes = classes1
		self.modelConfiguration = r"API\cfgs\YoloV3SIH(NoPlate-Final).cfg"
		# self.modelConfiguration = r"cfgs\YoloV3SIH(NoPlate-Final).cfg"
		self.modelWeights = r"API\weights\YoloV3SIH(NoPlate-Final)_final.weights"
		# self.modelWeights = r"weights\YoloV3SIH(NoPlate-Final)_final.weights"
		self.net = cv2.dnn.readNetFromDarknet(self.modelConfiguration, self.modelWeights)
		# self.net.setPreferableBackend(cv2.dnn.DNN_BACKEND_CUDA)
		# self.net.setPreferableTarget(cv2.dnn.DNN_TARGET_CUDA)
		self.net.setPreferableBackend(cv2.dnn.DNN_BACKEND_OPENCV)
		self.net.setPreferableTarget(cv2.dnn.DNN_TARGET_CPU)


	def Check_if_NoPlate_Exist(self,frame):
		nd1= time.time()
		self.blob = cv2.dnn.blobFromImage(frame, 1/255, (self.inpWidth, self.inpHeight), [0,0,0], 1, crop=False)
		self.net.setInput(self.blob)
		self.outs = self.net.forward(self.getOutputsNames())
		self.res = self.postprocess(frame)
		nd2= time.time()
		print('[NoPlate Detection Time]',nd2-nd1)
		retVal = 0
		if self.res != []:
			retVal = 1
		return (retVal,self.res)
		# t, _ = net.getPerfProfile()
		# label = 'Inference time: %.2f ms' % (t * 1000.0 / cv2.getTickFrequency())

	def getOutputsNames(self):
		self.layersNames = self.net.getLayerNames()
		return [self.layersNames[i[0] - 1] for i in self.net.getUnconnectedOutLayers()]
   
	def postprocess(self,frame):
		frameHeight = frame.shape[0]
		frameWidth = frame.shape[1]

		classIds = []
		confidences = []
		boxes = []
		for out in self.outs:
			print("out.shape : ", out.shape)
			for detection in out:
				scores = detection[5:]
				classId = np.argmax(scores)
				confidence = scores[classId]
				if detection[4]>self.confThreshold:
					#print(detection[4], " - ", scores[classId], " - th : ", confThreshold)
					#print(detection)
					pass
				if confidence > self.confThreshold:
					center_x = int(detection[0] * frameWidth)
					center_y = int(detection[1] * frameHeight)
					width = int(detection[2] * frameWidth)
					height = int(detection[3] * frameHeight)
					left = int(center_x - width / 2)
					top = int(center_y - height / 2)
					classIds.append(classId)
					confidences.append(float(confidence))
					boxes.append([left, top, width, height])

		left,width,top,height = '','','',''
		indices = cv2.dnn.NMSBoxes(boxes, confidences, self.confThreshold, self.nmsThreshold)
		if len(indices) > 1:
			print('Multiple Noplate in Image')
		noPlates = []
		for i in indices:
			i = i[0]
			box = boxes[i]
			left = box[0]
			top = box[1]
			width = box[2]
			height = box[3]
			print()
			noPlates.append((top,left,width,height))
			#l,t,r,b = self.convert(frame.shape,(left,right,top,bottom))
			#drawPred(classIds[i], confidences[i], left, top, left + width, top + height,file2)
		return noPlates

		def drawPred(self,classId, conf, left, top, right, bottom):
			cv2.rectangle(frame, (left, top), (right, bottom), (255, 178, 50), 2)
			label1 = '%.2f' % conf
			if classes:
				assert(classId < len(classes))
				label = '%s' % (classes[classId])
				print(forindex.index(label) ," : ", label1)
			labelSize, baseLine = cv2.getTextSize(label, cv2.FONT_HERSHEY_SIMPLEX, 0.5, 1)
			top = max(top, labelSize[1])
			l,t,r,b = convert(frame.shape,(left,right,top,bottom))
			cv2.putText(frame, label, (left, top), cv2.FONT_HERSHEY_SIMPLEX, 0.5, (0,0,255), 1)

		def convertToYolo(self, size, box):
			dw = 1./size[1]
			dh = 1./size[0]
			x = (box[0] + box[1])/2.0
			y = (box[2] + box[3])/2.0
			w = box[1] - box[0]
			h = box[3] - box[2]
			x = x*dw
			w = w*dw
			y = y*dh
			h = h*dh 
			return (x,y,w,h)

# document = Document()
# document.add_heading('NoPlate Detection Results Recognition', 0)
if __name__ == '__main__':
	no = 0
	NpdObj = No_Plate_Detector()
	lis = glob.glob(r'E:\0NewDev\SIH\Model_Test\TestImg\\*.jpg')
	font = cv2.FONT_HERSHEY_SIMPLEX
	ret = True
	for img in lis:
		if ret:
			frame = cv2.imread(img)
			ret = NpdObj.Check_if_NoPlate_Exist(frame)

			print(ret)
			if ret[0]:
				print(ret)
				ress = ret[1]
				for res in ress:
					print(res)
					no+=1
					cv2.rectangle(frame, (res[1], res[0]), (res[1]+res[2], res[0]+res[3]), (255, 178, 50), 7)
					cv2.imwrite(r'E:\0NewDev\SIH\Model_Test\TestImg\plate'+str(no)+'.jpg',frame[res[0]:res[0]+res[3],res[1]:res[1]+res[2]])
					document.add_picture(r'E:\0NewDev\SIH\Model_Test\TestImg\plate'+str(no)+'.jpg', width=Inches(1.25))
				
				cv2.imwrite(r'E:\0NewDev\SIH\Model_Test\TestImg\Frame'+str(no)+'.jpg',frame)
				document.add_picture(r'E:\0NewDev\SIH\Model_Test\TestImg\Frame'+str(no)+'.jpg', width=Inches(2.25))
				document.add_paragraph("________________________________________________________________________________")
	document.add_page_break()

	document.save('PlateDetaction.docx')